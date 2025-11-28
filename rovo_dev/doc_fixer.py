from __future__ import annotations
import argparse
from pathlib import Path
from datetime import datetime
import json
from docx import Document
from docx.shared import Pt
from .extractor import RovoExtractor
from .linker import RovoLinker
from .knowledge import KnowledgeBase
from .llm_client import LLMClient
from .semantic_mapper import SemanticMapper


def parse_args():
    ap = argparse.ArgumentParser(description="Generate normalized, corrected versions of documents using learned mappings")
    ap.add_argument("--folder", "-f", required=True, help="Folder to process (absolute or relative)")
    ap.add_argument("--adapter", type=str, default=None, help="Path to trained adapter .pt (optional)")
    ap.add_argument("--semantic-model", default="multi-qa-mpnet-base-dot-v1", help="SentenceTransformer model name")
    ap.add_argument("--semantic-threshold", type=float, default=0.4, help="Cosine threshold for mapping acceptance")
    ap.add_argument("--no-index-map", action="store_true", help="Disable index-based 1-to-1 mapping")
    ap.add_argument("--use-ai", action="store_true", help="Enable AI fallback (requires GEMINI_API_KEY)")
    ap.add_argument("--write-docx", action="store_true", help="Also write corrected DOCX (in addition to Markdown)")
    return ap.parse_args()


def find_project_folders(start_path: Path) -> list[Path]:
    def is_project(p: Path) -> bool:
        return (p / "STEP 2").exists() or (p / "Step 2").exists() or any(f.suffix.lower()==".docx" and not f.name.startswith("~$") for f in p.glob("*.docx"))

    if is_project(start_path):
        return [start_path]
    out = []
    for child in start_path.iterdir():
        if child.is_dir() and is_project(child):
            out.append(child)
    return out


def write_markdown(doc: dict, out_path: Path):
    lines = []
    lines.append(f"# {doc['lens']}")
    if doc.get('summary'):
        lines.append("\n## Summary\n")
        lines.append(doc['summary'])
    for p in sorted(doc.get('patterns', []), key=lambda x: x.get('pattern_number', 0)):
        lines.append(f"\n## Pattern {p.get('pattern_number')}: {p.get('title','')}")
        ov = p.get('overview') or ''
        if ov:
            lines.append(f"\n{ov}")
        vars = sorted(p.get('variations', []), key=lambda v: v.get('variation_number', 0))
        if vars:
            lines.append("\n### Variations\n")
        for v in vars:
            lines.append(f"- Variation {v.get('variation_number')}: {v.get('title','')}")
            content = v.get('content') or ''
            if content:
                lines.append(f"  \n  {content}")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")


def write_docx(doc: dict, out_path: Path):
    d = Document()
    styles = d.styles
    try:
        styles['Normal'].font.name = 'Calibri'
        styles['Normal'].font.size = Pt(11)
    except Exception:
        pass

    d.add_heading(doc['lens'], level=1)
    if doc.get('summary'):
        d.add_heading('Summary', level=2)
        for para in doc['summary'].split('\n'):
            d.add_paragraph(para)
    for p in sorted(doc.get('patterns', []), key=lambda x: x.get('pattern_number', 0)):
        d.add_heading(f"Pattern {p.get('pattern_number')}: {p.get('title','')}", level=2)
        ov = p.get('overview') or ''
        if ov:
            for para in ov.split('\n'):
                d.add_paragraph(para)
        vars = sorted(p.get('variations', []), key=lambda v: v.get('variation_number', 0))
        if vars:
            d.add_heading('Variations', level=3)
        for v in vars:
            d.add_paragraph(f"Variation {v.get('variation_number')}: {v.get('title','')}")
            content = v.get('content') or ''
            if content:
                for para in content.split('\n'):
                    d.add_paragraph(para)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(out_path))


def main():
    args = parse_args()
    start = Path(args.folder)
    projects = find_project_folders(start)
    if not projects:
        print(f"No projects found at {start}")
        return 1

    kb = KnowledgeBase()
    llm = LLMClient(enabled=args.use_ai)
    extractor = RovoExtractor()

    for proj in projects:
        print(f"Fixing and normalizing project: {proj}")
        step2 = proj / "STEP 2"
        if not step2.exists():
            step2 = proj / "Step 2"
        target = step2 if step2.exists() else proj

        semantic = SemanticMapper(model_name=args.semantic_model, threshold=args.semantic_threshold, adapter_path=args.adapter)
        linker = RovoLinker(kb=kb, llm=llm, semantic=semantic, allow_index=not args.no_index_map)

        normalized_docs = []
        for f in target.glob("*.docx"):
            if f.name.startswith("~$"):
                continue
            parsed = extractor.parse_document(f)
            if not parsed:
                continue
            linked_patterns = linker.link(parsed['lens'], parsed['patterns'], parsed['variations'])
            normalized_docs.append({
                "lens": parsed['lens'],
                "file_path": parsed['file_path'],
                "summary": parsed['summary'],
                "patterns": linked_patterns
            })

        if not normalized_docs:
            print(f"No documents processed for {proj.name}")
            continue

        # write outputs
        norm_dir = Path(__file__).resolve().parent / "normalized" / proj.name.lower()
        report_dir = Path(__file__).resolve().parent / "normalized" / proj.name.lower()
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        patch_report = {
            "project": proj.name,
            "timestamp": ts,
            "files": []
        }
        for nd in normalized_docs:
            stem = Path(nd['file_path']).stem
            md_out = norm_dir / f"{stem}_normalized.md"
            write_markdown(nd, md_out)
            if args.write_docx:
                docx_out = norm_dir / f"{stem}_normalized.docx"
                write_docx(nd, docx_out)
            # collect stats
            total_vars = sum(len(p.get('variations', [])) for p in nd.get('patterns', []))
            sem_count = sum(1 for p in nd.get('patterns', []) for v in p.get('variations', []) if v.get('_mapped_semantic'))
            fuzzy_count = sum(1 for p in nd.get('patterns', []) for v in p.get('variations', []) if v.get('_mapped_score') is not None and not v.get('_mapped_semantic'))
            fallback_count = sum(1 for p in nd.get('patterns', []) for v in p.get('variations', []) if not v.get('_mapped_semantic') and v.get('_mapped_score') is None)
            patch_report['files'].append({
                "file": nd['file_path'],
                "markdown": str(md_out),
                "docx": str(norm_dir / f"{stem}_normalized.docx") if args.write_docx else None,
                "patterns": len(nd.get('patterns', [])),
                "variations": total_vars,
                "mapped_semantic": sem_count,
                "mapped_fuzzy": fuzzy_count,
                "mapped_fallback": fallback_count,
            })

        report_path = report_dir / f"patch_report_{proj.name.lower()}_{ts}.json"
        report_dir.mkdir(parents=True, exist_ok=True)
        with open(report_path, "w", encoding="utf-8") as f:
            json.dump(patch_report, f, indent=2)
        print(f"Patch report saved: {report_path}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
