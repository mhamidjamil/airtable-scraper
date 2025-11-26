"""
Airtable Manager - Pattern Extraction Script
Extracts patterns, variations, lenses, and sources from .docx files

For AI Context: See CONTEXT_FOR_AI.md for complete project details
Known Issues: 
- Variations numbered 6-10 (not 1-5) - preserved as-is
- Pattern reference may not match actual pattern - logged for review
- Some implicit variations without explicit numbers - auto-numbered
"""

import os
import json
import re
import docx
from datetime import datetime
from pathlib import Path
from collections import defaultdict
from typing import List, Dict, Tuple, Optional


class ExtractionLogger:
    """
    Comprehensive logging for extraction process
    Tracks successes, warnings, errors, and data quality issues
    """
    
    def __init__(self, log_dir: str):
        self.log_dir = log_dir
        self.successful = []
        self.warnings = []
        self.errors = []
        self.skipped = []
        self.start_time = datetime.now()
        
        # Quality metrics
        self.total_patterns = 0
        self.total_variations = 0
        self.files_with_summary = 0
        self.variation_mismatches = []  # Track variation pattern_ref != actual pattern
    
    def log_success(self, file_path: str, pattern_count: int, variation_count: int, 
                   has_summary: bool):
        """Log successfully processed file"""
        self.successful.append({
            "file": file_path,
            "patterns": pattern_count,
            "variations": variation_count,
            "summary": has_summary
        })
        self.total_patterns += pattern_count
        self.total_variations += variation_count
        if has_summary:
            self.files_with_summary += 1
    
    def log_warning(self, file_path: str, warning_msg: str):
        """Log quality warning (non-fatal)"""
        self.warnings.append({
            "file": file_path,
            "warning": warning_msg
        })
    
    def log_error(self, file_path: str, error_msg: str):
        """Log processing error"""
        self.errors.append({
            "file": file_path,
            "error": error_msg
        })
    
    def log_skip(self, file_path: str, reason: str):
        """Log skipped file"""
        self.skipped.append({
            "file": file_path,
            "reason": reason
        })
    
    def log_variation_mismatch(self, file_path: str, variation_num: int, 
                              stated_pattern: int, actual_pattern: int):
        """
        Log when variation says it belongs to Pattern X but is assigned to Pattern Y
        This is informational - we still use document position for assignment
        """
        self.variation_mismatches.append({
            "file": file_path,
            "variation": variation_num,
            "stated_pattern": stated_pattern,
            "actual_pattern": actual_pattern
        })
    
    def save_log(self):
        """Save comprehensive log file"""
        timestamp = self.start_time.strftime("%Y%m%d_%H%M%S")
        log_path = os.path.join(self.log_dir, f"extraction_{timestamp}.log")
        
        with open(log_path, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("PATTERN EXTRACTION LOG\n")
            f.write("=" * 80 + "\n")
            f.write(f"Timestamp: {self.start_time.isoformat()}\n")
            f.write(f"Duration: {datetime.now() - self.start_time}\n")
            f.write("\n")
            
            # Summary Stats
            f.write("=" * 80 + "\n")
            f.write("SUMMARY STATISTICS\n")
            f.write("=" * 80 + "\n")
            total_files = len(self.successful) + len(self.skipped) + len(self.errors)
            f.write(f"Total files processed: {total_files}\n")
            f.write(f"✓ Successful: {len(self.successful)}\n")
            f.write(f"⊘ Skipped: {len(self.skipped)}\n")
            f.write(f"✗ Errors: {len(self.errors)}\n")
            f.write(f"⚠ Warnings: {len(self.warnings)}\n")
            f.write(f"\nData Extracted:\n")
            f.write(f"  - Total Patterns: {self.total_patterns}\n")
            f.write(f"  - Total Variations: {self.total_variations}\n")
            f.write(f"  - Files with Summary: {self.files_with_summary}\n")
            f.write("\n")
            
            # Successful files
            if self.successful:
                f.write("=" * 80 + "\n")
                f.write("SUCCESSFULLY PROCESSED FILES\n")
                f.write("=" * 80 + "\n")
                for item in self.successful:
                    f.write(f"✓ {item['file']}\n")
                    f.write(f"  Patterns: {item['patterns']}, ")
                    f.write(f"Variations: {item['variations']}, ")
                    f.write(f"Summary: {'Yes' if item['summary'] else 'No'}\n")
                f.write("\n")
            
            # Warnings
            if self.warnings:
                f.write("=" * 80 + "\n")
                f.write("QUALITY WARNINGS\n")
                f.write("=" * 80 + "\n")
                for item in self.warnings:
                    f.write(f"⚠ {item['file']}\n")
                    f.write(f"  {item['warning']}\n")
                f.write("\n")
            
            # Variation Mismatches
            if self.variation_mismatches:
                f.write("=" * 80 + "\n")
                f.write("VARIATION PATTERN REFERENCE MISMATCHES\n")
                f.write("(Informational - variations linked by position, not reference)\n")
                f.write("=" * 80 + "\n")
                for item in self.variation_mismatches:
                    f.write(f"⚠ {item['file']}\n")
                    f.write(f"  Variation {item['variation']}: ")
                    f.write(f"Says Pattern {item['stated_pattern']}, ")
                    f.write(f"Assigned to Pattern {item['actual_pattern']}\n")
                f.write("\n")
            
            # Skipped files
            if self.skipped:
                f.write("=" * 80 + "\n")
                f.write("SKIPPED FILES\n")
                f.write("=" * 80 + "\n")
                for item in self.skipped:
                    f.write(f"⊘ {item['file']}\n")
                    f.write(f"  Reason: {item['reason']}\n")
                f.write("\n")
            
            # Errors
            if self.errors:
                f.write("=" * 80 + "\n")
                f.write("ERRORS\n")
                f.write("=" * 80 + "\n")
                for item in self.errors:
                    f.write(f"✗ {item['file']}\n")
                    f.write(f"  Error: {item['error']}\n")
                f.write("\n")
        
        return log_path


def clean_label(text: str) -> str:
    """
    Remove common field labels from text content
    Labels: Explanation, Inner war / choice, Sources
    """
    if not text:
        return ""
    
    labels = [
        "Explanation:",
        "Inner war / choice:",
        "Sources:",
        "Explanation :",
        "Inner war / choice :",
        "Sources :"
    ]
    
    cleaned = text.strip()
    for label in labels:
        if cleaned.lower().startswith(label.lower()):
            cleaned = cleaned[len(label):].strip()
            break
    
    return cleaned


def clean_text(text: str) -> str:
    """
    Clean text by removing newlines and extra spaces
    Preserves content but makes it single-line for CSV compatibility
    """
    if not text:
        return ""
    
    # Replace literal \n with space
    text = text.replace('\\n', ' ')
    # Replace actual newlines with space
    text = text.replace('\n', ' ')
    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()


def extract_summary(paragraphs: List) -> Tuple[str, bool]:
    """
    Extract document summary (text before first Pattern)
    
    Rules:
    - Skip title (first line)
    - Collect paragraphs until "Pattern 1" or "Task 1"
    - Must have 2+ paragraphs OR 50+ characters
    
    Returns: (summary_text, has_valid_summary)
    """
    summary_lines = []
    first_line_skipped = False
    
    for para in paragraphs:
        text = para.text.strip()
        
        if not text:
            continue
        
        # Check if we've hit the pattern section
        if re.match(r'^(Task\s+1|TASK\s+1|Pattern\s+1|Part\s+I)', text, re.IGNORECASE):
            break
        
        # Skip title-only paragraphs (all caps, short)
        if text.isupper() and len(text) < 100:
            continue
        
        # Skip separator lines
        if re.match(r'^[_\-=]{3,}$', text):
            continue
        
        # Skip first meaningful line (document title)
        if not first_line_skipped:
            first_line_skipped = True
            continue
        
        summary_lines.append(text)
    
    # Join lines
    summary = "\n\n".join(summary_lines)
    
    # Validation: 2+ lines OR 50+ chars
    has_valid_summary = len(summary_lines) >= 2 or len(summary) > 50
    
    if has_valid_summary:
        summary = clean_text(summary)
        return summary, True
    
    return "", False


def extract_patterns(paragraphs: List) -> List[Dict]:
    """
    Extract patterns from document
    
    Pattern format:
    Pattern X: [Title]
    Explanation: [Overview]
    Inner war / choice: [Choice]
    Sources: [Source]
    
    Returns: List of pattern dictionaries
    """
    patterns = []
    i = 0
    
    while i < len(paragraphs):
        text = paragraphs[i].text.strip()
        
        # Check for pattern header: "Pattern 1: Title"
        pattern_match = re.match(r'^Pattern\s+(\d+):\s*(.+)$', text, re.IGNORECASE)
        
        if pattern_match:
            pattern_number = int(pattern_match.group(1))
            title = pattern_match.group(2).strip()
            
            # Collect next 3 non-empty paragraphs (overview, choice, source)
            overview = ""
            choice = ""
            source = ""
            
            j = i + 1
            section_index = 0
            
            while j < len(paragraphs) and section_index < 3:
                para_text = paragraphs[j].text.strip()
                
                if not para_text:
                    j += 1
                    continue
                
                # Stop if we hit another pattern or variation
                if re.match(r'^(Pattern|Variation)\s+\d+', para_text, re.IGNORECASE):
                    break
                
                # Clean label and text
                cleaned_text = clean_label(para_text)
                cleaned_text = clean_text(cleaned_text)
                
                # Assign to correct field
                if section_index == 0:
                    overview = cleaned_text
                elif section_index == 1:
                    choice = cleaned_text
                elif section_index == 2:
                    source = cleaned_text
                
                section_index += 1
                j += 1
            
            patterns.append({
                "pattern_number": pattern_number,
                "title": title,
                "overview": overview,
                "choice": choice,
                "source": source,
                "variations": []  # Will be populated later
            })
        
        i += 1
    
    return patterns


def extract_variations(paragraphs: List, logger: ExtractionLogger, 
                      file_path: str) -> List[Dict]:
    """
    Extract variations from document
    
    Handles multiple formats:
    1. "VARIATION 6 – PATTERN 1: Title"
    2. "– PATTERN 5: Title" (missing variation number)
    3. "Variation 6 — Title" (no PATTERN)
    4. "6 — Title" (just number, handle 0 as 10)
    5. "– UPPERCASE TITLE" (implicit variation)
    
    Returns: List of variation dictionaries with pattern_reference
    """
    variations = []
    i = 0
    current_var_num = 0  # For auto-numbering implicit variations
    
    while i < len(paragraphs):
        text = paragraphs[i].text.strip()
        
        variation_match = False
        variation_number = None
        pattern_ref = None
        title = None
        
        # Format 1: "VARIATION X – PATTERN Y: Title"
        match = re.match(r'^VARIATION\s+(\d+)\s*[–—-]\s*PATTERN\s+(\d+):\s*(.+)$', 
                        text, re.IGNORECASE)
        if match:
            variation_number = int(match.group(1))
            pattern_ref = int(match.group(2))
            title = match.group(3).strip()
            variation_match = True
            current_var_num = variation_number
        
        # Format 2: "– PATTERN X: Title"
        if not variation_match:
            match = re.match(r'^\s*[–—-]\s*PATTERN\s+(\d+):\s*(.+)$', 
                           text, re.IGNORECASE)
            if match:
                pattern_ref = int(match.group(1))
                title = match.group(2).strip()
                variation_number = pattern_ref  # Use pattern num as variation num
                variation_match = True
                current_var_num = variation_number
        
        # Format 3: "Variation X — Title"
        if not variation_match:
            match = re.match(r'^Variation\s+(\d+)\s*[–—-]\s*(.+)$', 
                           text, re.IGNORECASE)
            if match:
                variation_number = int(match.group(1))
                title = match.group(2).strip()
                pattern_ref = 1  # Default to pattern 1
                variation_match = True
                current_var_num = variation_number
        
        # Format 4: "X — Title" (handle 0 as 10)
        if not variation_match:
            match = re.match(r'^(\d+)\s*[–—-]\s*(.+)$', text)
            if match:
                num = int(match.group(1))
                if num == 0:
                    num = 10  # Known typo
                variation_number = num
                title = match.group(2).strip()
                pattern_ref = 1
                variation_match = True
                current_var_num = variation_number
        
        # Format 5: "– UPPERCASE TITLE" (implicit)
        if not variation_match:
            match = re.match(r'^\s*[–—-]\s*([A-Z\s]+)$', text)
            if match:
                title = match.group(1).strip()
                if len(title) > 3:  # Avoid noise
                    current_var_num += 1
                    variation_number = current_var_num
                    pattern_ref = 1
                    variation_match = True
        
        if variation_match:
            # Get next non-empty paragraph as content
            content = ""
            j = i + 1
            
            while j < len(paragraphs):
                para_text = paragraphs[j].text.strip()
                
                if not para_text:
                    j += 1
                    continue
                
                # Stop if we hit another pattern or variation
                stop_pattern = (
                    re.match(r'^(Pattern|Variation|VARIATION|\d+\s*[–—-])\s', 
                           para_text, re.IGNORECASE) or
                    re.match(r'^\s*[–—-]\s*[A-Z\s]+$', para_text)
                )
                if stop_pattern:
                    break
                
                content = clean_text(para_text)
                break
            
            if not content:
                logger.log_warning(file_path, 
                    f"Variation {variation_number} has no content")
            
            variations.append({
                "variation_number": variation_number,
                "pattern_reference": pattern_ref,
                "title": title,
                "content": content
            })
        
        i += 1
    
    return variations


def link_variations_to_patterns(patterns: List[Dict], variations: List[Dict],
                                logger: ExtractionLogger, file_path: str) -> List[Dict]:
    """
    Link variations to patterns
    
    Strategy:
    - Primary: Assign by document position (sequential)
    - Secondary: Check pattern_reference for validation
    - Log mismatches for review (informational only)
    
    Returns: Patterns with variations populated
    """
    # Create pattern map by number
    pattern_map = {p["pattern_number"]: p for p in patterns}
    
    for variation in variations:
        var_num = variation["variation_number"]
        stated_pattern = variation.get("pattern_reference", 1)
        
        # Assign to pattern by reference (if exists) or default to pattern 1
        if stated_pattern in pattern_map:
            actual_pattern = stated_pattern
        else:
            # Fallback: assign to first pattern
            actual_pattern = patterns[0]["pattern_number"] if patterns else None
            if actual_pattern:
                logger.log_warning(file_path, 
                    f"Variation {var_num} references Pattern {stated_pattern} " +
                    f"(not found), assigned to Pattern {actual_pattern}")
        
        if actual_pattern:
            # Add to pattern
            pattern_map[actual_pattern]["variations"].append({
                "variation_number": var_num,
                "title": variation["title"],
                "content": variation["content"]
            })
            
            # Log mismatch if any (informational)
            if stated_pattern != actual_pattern:
                logger.log_variation_mismatch(
                    file_path, var_num, stated_pattern, actual_pattern
                )
    
    return patterns


def process_metas_file(file_path: str, base_folder: str, 
                      logger: ExtractionLogger) -> Optional[Dict]:
    """
    Extract METAS data from METAS folder files
    
    Structure:
    - title: Filename (without .docx)
    - subtitle: First meaningful paragraph
    - content: Full text (all paragraphs)
    - base_folder: Parent folder name
    
    Returns: METAS dictionary or None if error
    """
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        if not paragraphs:
            logger.log_skip(file_path, "No content in METAS file")
            return None
        
        # Title from filename
        title = Path(file_path).stem
        
        # Subtitle from first paragraph
        subtitle = paragraphs[0] if paragraphs else ""
        
        # Content from all paragraphs
        content = "\n\n".join(paragraphs)
        content = clean_text(content)
        
        return {
            "title": title,
            "subtitle": subtitle,
            "content": content,
            "base_folder": base_folder
        }
        
    except Exception as e:
        logger.log_error(file_path, f"METAS extraction error: {str(e)}")
        return None


def process_pattern_file(file_path: str, base_folder: str, 
                        logger: ExtractionLogger) -> Optional[Dict]:
    """
    Process a single pattern document file
    
    Extracts:
    - Summary
    - Patterns (with overview, choice, source)
    - Variations (linked to patterns)
    - Metadata (lens from filename, base_folder from path)
    
    Returns: Document data dictionary or None if skipped/error
    """
    rel_path = os.path.relpath(file_path)
    
    try:
        # Read docx
        doc = docx.Document(file_path)
        paragraphs = doc.paragraphs
        
        # Extract components
        summary, has_summary = extract_summary(paragraphs)
        patterns = extract_patterns(paragraphs)
        variations = extract_variations(paragraphs, logger, rel_path)
        
        # Validation: Must have patterns
        if len(patterns) == 0:
            logger.log_skip(rel_path, "No patterns found")
            return None
        
        # Validation: Must have summary
        if not has_summary:
            logger.log_skip(rel_path, "No valid summary (need 2+ paragraphs or 50+ chars)")
            return None
        
        # Link variations to patterns
        patterns = link_variations_to_patterns(patterns, variations, logger, rel_path)
        
        # Extract metadata
        lens_name = Path(file_path).stem  # Filename without extension
        
        # Quality check: Warn if summary is unusually long
        if len(summary) > 2000:
            logger.log_warning(rel_path, f"Very long summary ({len(summary)} chars)")
        
        # Quality check: Warn if patterns have no variations
        patterns_without_vars = [p for p in patterns if len(p["variations"]) == 0]
        if patterns_without_vars:
            logger.log_warning(rel_path, 
                f"{len(patterns_without_vars)} patterns have no variations")
        
        # Log success
        total_variations = sum(len(p["variations"]) for p in patterns)
        logger.log_success(rel_path, len(patterns), total_variations, has_summary)
        
        return {
            "lens": lens_name,
            "base_folder": base_folder,
            "file_path": rel_path.replace(os.sep, '/'),
            "summary": summary,
            "patterns": patterns
        }
        
    except Exception as e:
        logger.log_error(rel_path, str(e))
        return None


def extract_folder(folder_path: str, output_dir: str, log_dir: str) -> Dict:
    """
    Extract all data from a folder (e.g., BIOME)
    
    Processes:
    - METAS files (from METAS subfolder)
    - Pattern files (from STEP 2 subfolder, or root if no STEP 2)
    
    Returns: Dictionary with extracted data
    """
    folder_path = Path(folder_path)
    base_folder_name = folder_path.name
    
    logger = ExtractionLogger(log_dir)
    
    print("=" * 80)
    print(f"EXTRACTING FOLDER: {base_folder_name}")
    print("=" * 80)
    print()
    
    # Data containers
    metas_data = []
    pattern_documents = []
    
    # Process METAS folder
    metas_folder = folder_path / "METAS"
    if metas_folder.exists() and metas_folder.is_dir():
        print(f"Processing METAS folder...")
        for file in metas_folder.glob("*.docx"):
            if file.name.startswith('~$'):
                continue
            print(f"  - {file.name}")
            meta_data = process_metas_file(str(file), base_folder_name, logger)
            if meta_data:
                metas_data.append(meta_data)
        print(f"  Extracted {len(metas_data)} METAS files\n")
    
    # Process pattern files (STEP 2 or root)
    step2_folder = None
    for subfolder in folder_path.iterdir():
        if subfolder.is_dir() and subfolder.name.lower() == "step 2":
            step2_folder = subfolder
            break
    
    pattern_folder = step2_folder if step2_folder else folder_path
    print(f"Processing pattern files from: {pattern_folder.name}/")
    
    pattern_files = list(pattern_folder.glob("*.docx"))
    pattern_files = [f for f in pattern_files if not f.name.startswith('~$')]
    
    for file in pattern_files:
        print(f"  - {file.name}")
        doc_data = process_pattern_file(str(file), base_folder_name, logger)
        if doc_data:
            pattern_documents.append(doc_data)
    
    print(f"  Extracted {len(pattern_documents)} pattern documents\n")
    
    # Save log
    log_path = logger.save_log()
    print(f"Log saved to: {log_path}\n")
    
    # Compile results
    results = {
        "base_folder": base_folder_name,
        "extraction_timestamp": datetime.now().isoformat(),
        "metas": metas_data,
        "documents": pattern_documents,
        "statistics": {
            "metas_count": len(metas_data),
            "documents_count": len(pattern_documents),
            "total_patterns": logger.total_patterns,
            "total_variations": logger.total_variations,
            "files_with_summary": logger.files_with_summary,
            "warnings": len(logger.warnings),
            "errors": len(logger.errors),
            "skipped": len(logger.skipped)
        }
    }
    
    # Save JSON output
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = os.path.join(output_dir, 
        f"{base_folder_name.lower()}_extracted_{timestamp}.json")
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    
    print(f"✓ Data saved to: {output_file}")
    print()
    
    # Print summary
    print("=" * 80)
    print("EXTRACTION SUMMARY")
    print("=" * 80)
    print(f"METAS extracted: {len(metas_data)}")
    print(f"Documents processed: {len(pattern_documents)}")
    print(f"Total patterns: {logger.total_patterns}")
    print(f"Total variations: {logger.total_variations}")
    print(f"Warnings: {len(logger.warnings)}")
    print(f"Errors: {len(logger.errors)}")
    print(f"Skipped: {len(logger.skipped)}")
    print()
    
    return results


def main():
    """Main execution"""
    # Configuration
    SOURCE_FOLDER = r"E:\Work\shoaib\upwork\new_extractions\BIOME"
    OUTPUT_DIR = r"E:\Work\shoaib\upwork\airtable_manager\data_output"
    LOG_DIR = r"E:\Work\shoaib\upwork\airtable_manager\logs"
    
    # Create directories
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(LOG_DIR, exist_ok=True)
    
    # Extract
    results = extract_folder(SOURCE_FOLDER, OUTPUT_DIR, LOG_DIR)
    
    print("=" * 80)
    print("EXTRACTION COMPLETE!")
    print("=" * 80)
    print()
    print("Next steps:")
    print("1. Review extraction log in logs/ folder")
    print("2. Review extracted JSON in data_output/ folder")
    print("3. Run prepare_airtable_data.py to generate CSV files")
    print()


if __name__ == "__main__":
    main()
