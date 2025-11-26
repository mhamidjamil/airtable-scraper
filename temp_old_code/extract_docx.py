from docx import Document
import sys

def extract_text_from_docx(docx_path, output_path):
    doc = Document(docx_path)
    
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    text = '\n'.join(full_text)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    
    print(f"Extracted text saved to {output_path}")

if __name__ == "__main__":
    extract_text_from_docx(
        r"E:\Work\shoaib\upwork\workflow-martin.docx",
        r"E:\Work\shoaib\upwork\workflow-martin.txt"
    )
