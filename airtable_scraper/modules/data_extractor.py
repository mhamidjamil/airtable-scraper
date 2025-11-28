import os
import re
import docx
import json
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from config import settings
from extraction_rules import VariationExtractor, SourceExtractor

class DataExtractor:
    def __init__(self, log_handler=None):
        self.logger = log_handler
        # Initialize extraction rule engines
        self.variation_extractor = VariationExtractor()
        self.source_extractor = SourceExtractor()
        
    def log(self, msg, level="info"):
        if self.logger:
            if level == "error":
                self.logger.error(msg)
            elif level == "warning":
                self.logger.warning(msg)
            else:
                self.logger.info(msg)
        else:
            print(f"[{level.upper()}] {msg}")

    def clean_text(self, text: str) -> str:
        """Clean text using extraction rules"""
        return self.variation_extractor.clean_text(text)

    def clean_label(self, text: str) -> str:
        """Remove common labels from text content using extraction rules"""
        return self.source_extractor.clean_label(text)
    
    def parse_sources(self, source_text: str, lens_name: str, base_folder: str) -> List[Dict]:
        """Parse multiple sources from source text using extraction rules"""
        return self.source_extractor.parse_sources(source_text, lens_name, base_folder, self.logger)

    def extract_summary(self, paragraphs: List) -> Tuple[str, bool]:
        summary_lines = []
        first_line_skipped = False
        
        for para in paragraphs:
            text = para.text.strip()
            if not text: continue
            
            # Stop at pattern start
            if re.match(r'^(Task\s+1|TASK\s+1|Pattern\s+1|Part\s+I)', text, re.IGNORECASE):
                break
            
            # Skip title/separators
            if (text.isupper() and len(text) < 100) or re.match(r'^[_\-=]{3,}$', text):
                continue
                
            if not first_line_skipped:
                first_line_skipped = True
                continue
                
            summary_lines.append(text)
        
        summary = "\n\n".join(summary_lines)
        has_valid = len(summary_lines) >= 2 or len(summary) > 50
        return (self.clean_text(summary), True) if has_valid else ("", False)

    # a: Pattern Extractor
    def extract_patterns(self, paragraphs: List) -> List[Dict]:
        patterns = []
        i = 0
        while i < len(paragraphs):
            text = paragraphs[i].text.strip()
            match = re.match(r'^Pattern\s+(\d+):\s*(.+)$', text, re.IGNORECASE)
            
            if match:
                p_num = int(match.group(1))
                title = match.group(2).strip()
                overview, choice, source = "", "", ""
                
                # Collect next 3 sections
                j = i + 1
                section_idx = 0
                while j < len(paragraphs) and section_idx < 3:
                    p_text = paragraphs[j].text.strip()
                    if not p_text:
                        j += 1
                        continue
                    
                    if re.match(r'^(Pattern|Variation)\s+\d+', p_text, re.IGNORECASE):
                        break
                        
                    cleaned = self.clean_text(self.clean_label(p_text))
                    if section_idx == 0: overview = cleaned
                    elif section_idx == 1: choice = cleaned
                    elif section_idx == 2: source = cleaned
                    
                    section_idx += 1
                    j += 1
                
                patterns.append({
                    "pattern_number": p_num,
                    "title": title,
                    "overview": overview,
                    "choice": choice,
                    "source": source,  # Keep original for backward compatibility
                    "variations": []
                })
            i += 1
        return patterns

    # c: Variation Extractor
    def extract_variations(self, paragraphs: List, file_path: str) -> List[Dict]:
        """Extract variations from document paragraphs using extraction rules"""
        return self.variation_extractor.extract_variations(paragraphs, file_path, self.logger)

    # e: Metas Extractor
    def extract_metas(self, file_path: str, base_folder: str) -> Optional[Dict]:
        try:
            doc = docx.Document(file_path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if not paras:
                return None
            
            # Look for title:subtitle pattern in first paragraph
            title = ""
            subtitle = ""
            content_start_idx = 0
            
            if paras:
                first_para = paras[0].strip()
                
                # Check if first paragraph contains title:subtitle pattern
                if ": " in first_para:
                    # Split on first colon to get title and subtitle
                    parts = first_para.split(": ", 1)
                    if len(parts) == 2:
                        title = parts[0].strip()
                        subtitle = parts[1].strip()
                        content_start_idx = 1
                    else:
                        # Fallback to filename as title, first para as subtitle
                        title = Path(file_path).stem
                        subtitle = first_para
                        content_start_idx = 1
                else:
                    # No colon pattern, use filename as title, first para as subtitle
                    title = Path(file_path).stem
                    subtitle = first_para
                    content_start_idx = 1
            
            # Content: everything after subtitle
            content_paras = paras[content_start_idx:]
            content = self.clean_text("\n\n".join(content_paras))
            
            return {
                "title": title,
                "subtitle": subtitle,
                "content": content,
                "base_folder": base_folder
            }
        except Exception as e:
            self.log(f"Error extracting META {file_path}: {str(e)}", "error")
            return None

    # Main Processing Method
    def process_folder(self, folder_input: str, extract_types: List[str] = None) -> Dict:
        # Handle both absolute paths and relative names
        if os.path.isabs(folder_input):
            folder_path = Path(folder_input)
        else:
            folder_path = settings.SOURCE_DIR / folder_input

        if not folder_path.exists():
            self.log(f"Folder not found: {folder_path}", "error")
            return {}

        folder_name = folder_path.name
        self.log(f"Processing folder: {folder_name} (Path: {folder_path})")
        
        extracted_data = {
            "base_folder": folder_name,
            "timestamp": datetime.now().isoformat(),
            "metas": [],
            "documents": []
        }

        # Determine what to extract
        if extract_types is None:
            extract_types = ['metas', 'lenses', 'sources', 'patterns', 'variations']
            
        should_extract_metas = 'metas' in extract_types
        should_extract_docs = any(t in extract_types for t in ['lenses', 'sources', 'patterns', 'variations'])

        # 1. Extract METAS
        if should_extract_metas:
            metas_dir = folder_path / "METAS"
            if metas_dir.exists():
                self.log(f"Found METAS directory: {metas_dir}")
                for f in metas_dir.glob("*.docx"):
                    if f.name.startswith("~$"): continue
                    meta = self.extract_metas(str(f), folder_name)
                    if meta: 
                        extracted_data["metas"].append(meta)
                        self.log(f"Extracted Meta: {meta['title']}")
            else:
                self.log(f"No METAS directory found in {folder_path}", "warning")

        # 2. Extract Documents (Patterns, Sources, Lenses, Variations)
        if should_extract_docs:
            step2_dir = folder_path / "STEP 2"
            if not step2_dir.exists():
                step2_dir = folder_path / "Step 2"
                
            target_dir = step2_dir if step2_dir.exists() else folder_path
            
            for f in target_dir.glob("*.docx"):
                if f.name.startswith("~$"): continue
                
                try:
                    doc = docx.Document(str(f))
                    paras = doc.paragraphs
                    
                    # Extract components
                    summary, has_summary = self.extract_summary(paras)
                    patterns = self.extract_patterns(paras)
                    variations = self.extract_variations(paras, f.name)
                    
                    if not patterns or not has_summary:
                        self.log(f"Skipping {f.name}: Missing patterns or summary", "warning")
                        continue

                    # Link variations to patterns
                    has_explicit_ref = any(v["pattern_reference"] != 1 for v in variations)
                    force_explicit_mapping = any(v.get("forced_explicit_mapping", False) for v in variations)
                    
                    if has_explicit_ref or force_explicit_mapping:
                        if force_explicit_mapping:
                            self.log(f"One-per-pattern mapping forced in {f.name}. Using 1-to-1 mapping.")
                        else:
                            self.log(f"Explicit pattern references detected in {f.name}. Using 1-to-1 mapping.")
                        p_map = {p["pattern_number"]: p for p in patterns}
                        
                        for v in variations:
                            target_pattern_num = v["variation_number"]
                            target = p_map.get(target_pattern_num)
                            
                            if target:
                                target["variations"].append({
                                    "variation_number": v["variation_number"],
                                    "title": v["title"],
                                    "content": v["content"]
                                })
                                self.log(f"Linked Variation {v['variation_number']} -> Pattern {target_pattern_num}")
                            else:
                                # Fallback: Link to Pattern 1 if no matching pattern number
                                if patterns:
                                    fallback = patterns[0]
                                    fallback["variations"].append({
                                        "variation_number": v["variation_number"],
                                        "title": v["title"],
                                        "content": v["content"]
                                    })
                                    self.log(f"Warning: Variation {v['variation_number']} has no matching Pattern {target_pattern_num}. Linked to Pattern 1 as fallback.", "warning")
                    else:
                        self.log(f"No explicit pattern references in {f.name}. Linking ALL to Pattern 1.")
                        if patterns:
                            target = patterns[0]
                            for v in variations:
                                target["variations"].append({
                                    "variation_number": v["variation_number"],
                                    "title": v["title"],
                                    "content": v["content"]
                                })
                            self.log(f"Linked {len(variations)} variations to Pattern 1: {target['title'][:30]}...")

                    # d: Lens Extractor
                    lens_name = f.stem
                    
                    # b: Source Extractor
                    all_sources = []
                    for pattern in patterns:
                        pattern_sources = self.parse_sources(pattern.get("source", ""), lens_name, folder_name)
                        all_sources.extend(pattern_sources)
                        pattern["parsed_sources"] = pattern_sources

                    extracted_data["documents"].append({
                        "lens": lens_name,
                        "base_folder": folder_name,
                        "file_path": str(f),
                        "summary": summary,
                        "patterns": patterns,
                        "sources": all_sources
                    })
                    
                except Exception as e:
                    self.log(f"Error processing {f.name}: {str(e)}", "error")

        # Save to JSON
        out_file = settings.DATA_DIR / f"{folder_name.lower()}_data.json"
        with open(out_file, 'w', encoding='utf-8') as f:
            json.dump(extracted_data, f, indent=2, ensure_ascii=False)
        
        self.log(f"Extraction complete. Saved to {out_file}")
        
        # Log extraction summary
        doc_count = len(extracted_data.get("documents", []))
        meta_count = len(extracted_data.get("metas", []))
        total_patterns = sum(len(doc.get("patterns", [])) for doc in extracted_data.get("documents", []))
        total_variations = sum(len(pattern.get("variations", [])) 
                          for doc in extracted_data.get("documents", []) 
                          for pattern in doc.get("patterns", []))

        self.log(f"Extraction Summary - Documents: {doc_count}, Patterns: {total_patterns}, Variations: {total_variations}, Metas: {meta_count}")

        return extracted_data
