import docx
import glob
import os

# Find the file
files = glob.glob(r'E:\Work\shoaib\upwork\new_extractions\BULLSHIT\STEP 2\*.docx')
target_file = None
for f in files:
    if "Garden" in f:
        target_file = f
        break

if not target_file:
    print("File not found")
    # Print all files to debug
    print("Available files:")
    for f in files:
        print(f)
    exit()

f = target_file
print(f"Reading {f}...")

doc = docx.Document(f)
print(f"Total paragraphs: {len(doc.paragraphs)}")

# Print paragraphs 50-100
for i, p in enumerate(doc.paragraphs):
    if 50 <= i <= 100:
        text = p.text.strip()
        if text:
            print(f"Line {i}: {text[:100]}")
